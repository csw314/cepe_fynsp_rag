"""Complete synthetic dashboard, RAG, ontology, landing, and report acceptance test."""

from __future__ import annotations

import json
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from pydantic import ValidationError

from cepe_fynsp.agents.rag_agent import answer_question, load_validated_rag_corpus
from cepe_fynsp.reporting.generator import REPORT_SECTIONS
from cepe_fynsp.schemas import DashboardManifest, DashboardQuestionPayload, RagRecord
from scripts.build_synthetic_ci import run_synthetic_build
from scripts.validate_static import validate_static


EXPECTED_VISUALIZATIONS = {
    "dashboard_01_pit_production": (
        "stacked_bar",
        "ranked_horizontal_bar",
        "ranked_horizontal_bar",
        "pareto_ranked_bar",
        "quality_scorecard_and_table",
        "reconciliation_variance_table",
    ),
    "dashboard_02_acquisition_schedule": (
        "stacked_bar",
        "ranked_table",
        "schedule_table",
        "exception_table",
        "site_year_heatmap_table",
        "bubble_plot",
    ),
    "dashboard_03_site_capacity": (
        "ranked_bar_table",
        "site_year_heatmap_table",
        "above_baseline_dependency_table",
        "suboffice_site_matrix",
        "yoy_change_table",
        "scope_quality_scorecard",
    ),
    "dashboard_04_priority_challenge": (
        "pareto_ranked_bar",
        "tier_funding_matrix",
        "priority_uniqueness_matrix",
        "classified_ranked_table",
        "traceability_scorecard",
        "account_integrator_completeness_table",
    ),
    "dashboard_05_findings_report_generator": (
        "finding_cards",
        "coverage_matrix",
        "risk_opportunity_heatmap_table",
        "exhibit_gallery_table",
        "citation_lineage_table",
        "report_outline_status",
    ),
}


def _lineage_values(value: object) -> set[str]:
    found: set[str] = set()
    if isinstance(value, dict):
        for key, nested in value.items():
            if key in {"source_record_id_sample", "source_row_id_sample"} and isinstance(
                nested, list
            ):
                found.update(str(item) for item in nested)
            else:
                found.update(_lineage_values(nested))
    elif isinstance(value, list):
        for nested in value:
            found.update(_lineage_values(nested))
    return found


def test_complete_synthetic_build(tmp_path: Path) -> None:
    counts = run_synthetic_build(tmp_path)
    assert counts["dashboards"] == 5
    assert counts["payloads"] == 30
    assert counts["rag_records"] == 30

    payload_root = tmp_path / "data" / "curated" / "dashboard_payloads"
    for manifest_path in sorted(payload_root.glob("dashboard_*/manifest.json")):
        manifest = DashboardManifest.model_validate_json(manifest_path.read_text(encoding="utf-8"))
        assert [entry.question_id for entry in manifest.payloads] == [
            f"q{index}" for index in range(1, 7)
        ]
        observed_types: list[str] = []
        payload_lineage: set[str] = set()
        for entry in manifest.payloads:
            payload = DashboardQuestionPayload.model_validate_json(
                (manifest_path.parent / entry.file).read_text(encoding="utf-8")
            )
            if manifest.dashboard_id == "dashboard_01_pit_production" and entry.question_id == "q1":
                unsupported = payload.model_dump(mode="json")
                unsupported["schema_version"] = "1.0"
                with pytest.raises(ValidationError, match="Unsupported dashboard payload"):
                    DashboardQuestionPayload.model_validate(unsupported)
            observed_types.append(payload.visualization.type)
            assert payload.active_filter_state["program_int_area"] == "Pit Production"
            assert payload.quality_summary.financial_completeness["aggregate_status"]
            assert all(metric.aggregate_status for metric in payload.metrics)
            assert payload.columns
            assert payload.ontology_references
            assert all(record.origin != "ai_generated_narrative" for record in payload.narrative)
            assert all(record.citations for record in payload.narrative)
            payload_lineage.update(_lineage_values(payload.lineage))
            if not (
                manifest.dashboard_id == "dashboard_01_pit_production"
                and entry.question_id in {"q5", "q6"}
            ):
                submission = payload.source_submission_type
                assert not (
                    "Federal Crosscuts" in submission and "Federal Site Splits" in submission
                )
        assert tuple(observed_types) == EXPECTED_VISUALIZATIONS[manifest.dashboard_id]
        rag_records = [
            RagRecord.model_validate_json(line)
            for line in (tmp_path / manifest.rag_context_file)
            .read_text(encoding="utf-8")
            .splitlines()
        ]
        assert all(record.payload_ids and record.citation_labels for record in rag_records)
        assert all(set(record.lineage_ids) <= payload_lineage for record in rag_records)

    landing = json.loads((payload_root / "landing_summary.json").read_text(encoding="utf-8"))
    assert len(landing["metrics"]) >= 6
    assert all(metric["payload_id"] for metric in landing["metrics"])
    assert not any(str(metric["display"]).startswith("$0") for metric in landing["metrics"][:2])

    report_manifest_path = (
        tmp_path / "data" / "reports" / "html" / "dashboard_05_report_manifest.json"
    )
    report_manifest = json.loads(report_manifest_path.read_text(encoding="utf-8"))
    assert tuple(report_manifest["sections"]) == REPORT_SECTIONS
    assert report_manifest["ai_assisted_text"] is False
    assert len(report_manifest["exhibits"]) == 7
    assert report_manifest["citation_count"] >= 30
    docx_path = tmp_path / report_manifest["docx_file"]
    document = Document(docx_path)
    heading_text = {paragraph.text for paragraph in document.paragraphs}
    assert set(REPORT_SECTIONS) <= heading_text
    assert any(len(table.rows) == 31 for table in document.tables)
    with ZipFile(docx_path) as archive:
        assert len([name for name in archive.namelist() if name.startswith("word/media/")]) == 7
    citation_manifest = json.loads(
        (tmp_path / report_manifest["citation_manifest_file"]).read_text(encoding="utf-8")
    )
    assert all(
        citation["chart_id"] and citation["citation_id"] and citation["metric_definition"]
        for citation in citation_manifest["citations"]
    )
    assert any(citation["finding_ids"] for citation in citation_manifest["citations"])
    combined_claims = " ".join(
        citation["claim_text"] for citation in citation_manifest["citations"]
    ).lower()
    assert "guaranteed executable" not in combined_claims

    corpus = load_validated_rag_corpus(tmp_path)
    answer = answer_question(
        corpus,
        "How much funding is programmed and what are the limitations?",
        filter_state={"program_int_area": "Pit Production"},
    )
    assert answer.status == "answered"
    assert answer.citation_labels and answer.payload_ids and answer.ontology_ids
    assert answer.interpretations == ()
    assert answer.ai_generated_narrative is None

    first_manifest_path = next(payload_root.glob("dashboard_*/manifest.json"))
    first_manifest = DashboardManifest.model_validate_json(
        first_manifest_path.read_text(encoding="utf-8")
    )
    rag_path = tmp_path / first_manifest.rag_context_file
    lines = rag_path.read_text(encoding="utf-8").splitlines()
    invalid = json.loads(lines[0])
    invalid["ontology_ids"] = ["missing:ontology:node"]
    lines[0] = json.dumps(invalid)
    rag_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    with pytest.raises(ValueError, match="Invalid RAG references"):
        load_validated_rag_corpus(tmp_path)


def test_static_frontend_is_offline_and_complete() -> None:
    validate_static(Path(__file__).resolve().parents[2])
