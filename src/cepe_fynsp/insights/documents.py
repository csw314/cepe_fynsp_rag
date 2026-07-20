"""Approval-gated guidance parsing, chunking, indexing, and local lexical retrieval."""

from __future__ import annotations

import hashlib
import json
import os
import re
import zipfile
from datetime import UTC, datetime
from pathlib import Path
from typing import Literal
from xml.etree import ElementTree

import yaml
from docx import Document
from pydantic import Field
from pypdf import PdfReader

from cepe_fynsp.insights.schemas import DocumentChunk, StrictInsightSchema

PARSER_VERSION = "guidance_parser_v2"
SUPPORTED_SUFFIXES = {".pdf", ".docx", ".pptx", ".txt", ".md", ".markdown"}
TOKEN_PATTERN = re.compile(r"[a-z0-9]+")
PPTX_SLIDE_PATTERN = re.compile(r"ppt/slides/slide([1-9][0-9]*)\.xml")
MAX_PPTX_SLIDES = 500
MAX_PPTX_SLIDE_XML_BYTES = 5_000_000
MAX_PPTX_TOTAL_SLIDE_XML_BYTES = 50_000_000


class ApprovedDocument(StrictInsightSchema):
    """One explicitly reviewed document eligible for local indexing."""

    path: str = Field(min_length=1, max_length=500)
    title: str = Field(min_length=1, max_length=300)
    document_type: Literal["guidance", "policy", "reference"] = "guidance"
    approved_for_asksage: bool = False
    classification: str = Field(default="internal", min_length=1, max_length=100)


class DocumentApprovalManifest(StrictInsightSchema):
    """Explicit allowlist; files absent from this manifest are never indexed."""

    schema_version: Literal["1.0"]
    documents: tuple[ApprovedDocument, ...] = ()


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _normalize_text(value: str) -> str:
    return re.sub(r"[ \t]+", " ", value.replace("\x00", "")).strip()


def _parse_pptx(path: Path) -> list[tuple[int | None, str | None, str]]:
    """Extract ordered visible slide text from a bounded local PPTX archive."""
    with zipfile.ZipFile(path) as archive:
        slides: list[tuple[int, zipfile.ZipInfo]] = []
        for member in archive.infolist():
            match = PPTX_SLIDE_PATTERN.fullmatch(member.filename)
            if match:
                slides.append((int(match.group(1)), member))
        slides.sort(key=lambda item: item[0])
        if len(slides) > MAX_PPTX_SLIDES:
            raise ValueError(f"PPTX exceeds the {MAX_PPTX_SLIDES}-slide limit.")
        if sum(member.file_size for _, member in slides) > MAX_PPTX_TOTAL_SLIDE_XML_BYTES:
            raise ValueError("PPTX slide XML exceeds the total extraction limit.")

        extracted: list[tuple[int | None, str | None, str]] = []
        for slide_number, member in slides:
            if member.file_size > MAX_PPTX_SLIDE_XML_BYTES:
                raise ValueError(f"PPTX slide {slide_number} exceeds the extraction limit.")
            content = archive.read(member)
            upper_content = content.upper()
            if b"<!DOCTYPE" in upper_content or b"<!ENTITY" in upper_content:
                raise ValueError(
                    f"PPTX slide {slide_number} contains unsupported XML declarations."
                )
            root = ElementTree.fromstring(content)
            text = _normalize_text(
                "\n".join(
                    element.text or ""
                    for element in root.iter()
                    if element.tag.endswith("}t") and (element.text or "").strip()
                )
            )
            if text:
                extracted.append((slide_number, f"Slide {slide_number}", text))
        return extracted


def parse_document(path: Path) -> list[tuple[int | None, str | None, str]]:
    """Extract page/section-aware text from a supported approved document."""
    suffix = path.suffix.casefold()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(f"Unsupported approved document type: {suffix or '<none>'}")
    if suffix == ".pdf":
        pages: list[tuple[int | None, str | None, str]] = []
        for page_number, page in enumerate(PdfReader(str(path)).pages, start=1):
            text = _normalize_text(page.extract_text() or "")
            if text:
                pages.append((page_number, None, text))
        return pages
    if suffix == ".docx":
        text = "\n".join(
            paragraph.text for paragraph in Document(str(path)).paragraphs if paragraph.text.strip()
        )
        normalized = _normalize_text(text)
        return [(None, None, normalized)] if normalized else []
    if suffix == ".pptx":
        return _parse_pptx(path)
    text = path.read_text(encoding="utf-8-sig")
    if suffix in {".md", ".markdown"}:
        sections: list[tuple[int | None, str | None, str]] = []
        heading: str | None = None
        buffer: list[str] = []
        for line in text.splitlines():
            if line.lstrip().startswith("#"):
                normalized = _normalize_text("\n".join(buffer))
                if normalized:
                    sections.append((None, heading, normalized))
                heading = line.lstrip("#").strip() or None
                buffer = []
            else:
                buffer.append(line)
        normalized = _normalize_text("\n".join(buffer))
        if normalized:
            sections.append((None, heading, normalized))
        return sections
    normalized = _normalize_text(text)
    return [(None, None, normalized)] if normalized else []


def _split_text(text: str, *, target_chars: int = 1600, overlap_chars: int = 200) -> list[str]:
    """Split text deterministically on paragraph/sentence boundaries with bounded overlap."""
    if len(text) <= target_chars:
        return [text]
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + target_chars)
        if end < len(text):
            boundary = max(text.rfind(". ", start, end), text.rfind("\n", start, end))
            if boundary > start + target_chars // 2:
                end = boundary + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(start + 1, end - overlap_chars)
    return chunks


def _resolve_approved_path(raw_docs_root: Path, configured: str) -> Path:
    candidate = (raw_docs_root / configured).resolve()
    try:
        candidate.relative_to(raw_docs_root.resolve())
    except ValueError as exc:
        raise ValueError("Approved document path escapes data/raw/docs.") from exc
    if not candidate.is_file():
        raise FileNotFoundError(f"Approved document is missing: {configured}")
    return candidate


def index_approved_documents(
    project_root: Path,
    *,
    approval_manifest: Path | None = None,
    output_path: Path | None = None,
) -> tuple[DocumentChunk, ...]:
    """Build a derived JSONL index from explicitly approved immutable documents."""
    root = project_root.resolve()
    manifest_path = approval_manifest or root / "config" / "approved_guidance_docs.yaml"
    manifest = DocumentApprovalManifest.model_validate(
        yaml.safe_load(manifest_path.read_text(encoding="utf-8")) or {}
    )
    raw_docs_root = root / "data" / "raw" / "docs"
    generated_at = datetime.now(UTC).isoformat()
    chunks: list[DocumentChunk] = []
    seen_paths: set[Path] = set()
    for approved in manifest.documents:
        if not approved.approved_for_asksage:
            continue
        path = _resolve_approved_path(raw_docs_root, approved.path)
        if path in seen_paths:
            raise ValueError("Approved document manifest contains a duplicate path.")
        seen_paths.add(path)
        source_hash = _sha256(path)
        source_file_id = f"document:{source_hash[:16]}"
        chunk_index = 0
        for page_number, section_heading, text in parse_document(path):
            for chunk_text in _split_text(text):
                identity = json.dumps(
                    [source_hash, page_number, section_heading, chunk_index, chunk_text],
                    ensure_ascii=False,
                    separators=(",", ":"),
                )
                chunk_id = f"guidance:{hashlib.sha256(identity.encode('utf-8')).hexdigest()[:20]}"
                chunks.append(
                    DocumentChunk(
                        chunk_id=chunk_id,
                        source_file_id=source_file_id,
                        source_file_hash=source_hash,
                        document_title=approved.title,
                        document_type=approved.document_type,
                        page_number=page_number,
                        section_heading=section_heading,
                        chunk_index=chunk_index,
                        chunk_text=chunk_text,
                        classification_metadata={
                            "classification": approved.classification,
                            "approval": "approved_for_asksage",
                        },
                        generated_at=generated_at,
                        parser_version=PARSER_VERSION,
                    )
                )
                chunk_index += 1
    destination = output_path or root / "data" / "curated" / "guidance_chunks" / "index.jsonl"
    destination.parent.mkdir(parents=True, exist_ok=True)
    temporary = destination.with_suffix(destination.suffix + ".tmp")
    with temporary.open("w", encoding="utf-8") as handle:
        for chunk in chunks:
            handle.write(chunk.model_dump_json() + "\n")
    os.replace(temporary, destination)
    return tuple(chunks)


def load_document_index(path: Path) -> tuple[DocumentChunk, ...]:
    """Load and validate a derived approved-document index without logging its text."""
    if not path.is_file():
        return ()
    return tuple(
        DocumentChunk.model_validate_json(line)
        for line in path.read_text(encoding="utf-8").splitlines()
        if line.strip()
    )


def retrieve_document_chunks(
    chunks: tuple[DocumentChunk, ...],
    query_text: str,
    *,
    limit: int = 5,
    max_characters: int = 12_000,
) -> tuple[DocumentChunk, ...]:
    """Return bounded lexical matches with stable deterministic ordering and citations."""
    query_tokens = set(TOKEN_PATTERN.findall(query_text.casefold()))
    ranked: list[tuple[int, str, DocumentChunk]] = []
    for chunk in chunks:
        evidence_tokens = set(
            TOKEN_PATTERN.findall(
                f"{chunk.document_title} {chunk.section_heading or ''} {chunk.chunk_text}".casefold()
            )
        )
        score = len(query_tokens & evidence_tokens)
        if score:
            ranked.append((score, chunk.chunk_id, chunk))
    ranked.sort(key=lambda item: (-item[0], item[1]))
    selected: list[DocumentChunk] = []
    used = 0
    for _, _, chunk in ranked:
        if len(selected) >= max(0, limit) or used + len(chunk.chunk_text) > max_characters:
            break
        selected.append(chunk)
        used += len(chunk.chunk_text)
    return tuple(selected)
