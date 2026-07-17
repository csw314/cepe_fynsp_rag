"""Deterministic Word-compatible management report generation."""

from __future__ import annotations

import html
import os
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw

from cepe_fynsp.config import load_settings
from cepe_fynsp.dashboards.dashboard_support import format_dollars, write_json
from cepe_fynsp.schemas import CitationRecord, DashboardManifest, DashboardQuestionPayload

REPORT_ID = "cepe_pit_production_fynsp_review_draft"
REPORT_SECTIONS = (
    "Executive Summary",
    "Review Scope and Data Sources",
    "Dashboard Question Answers",
    "Accuracy Findings",
    "Thoroughness Findings",
    "Uncertainty, Risk, and Opportunity Findings",
    "Data Limitations and Recommended Follow-Up",
    "Appendix: Source Lineage and Guidance Citations",
)


def _load_payloads(root: Path) -> dict[str, dict[str, DashboardQuestionPayload]]:
    """Load and validate all dashboard outputs used as report evidence."""
    dashboards: dict[str, dict[str, DashboardQuestionPayload]] = {}
    payload_root = root / "data" / "curated" / "dashboard_payloads"
    for manifest_path in sorted(payload_root.glob("*/manifest.json")):
        manifest = DashboardManifest.model_validate_json(manifest_path.read_text(encoding="utf-8"))
        dashboards[manifest.dashboard_id] = {
            entry.question_id: DashboardQuestionPayload.model_validate_json(
                (manifest_path.parent / entry.file).read_text(encoding="utf-8")
            )
            for entry in manifest.payloads
        }
    required = {
        "dashboard_01_pit_production",
        "dashboard_02_acquisition_schedule",
        "dashboard_03_site_capacity",
        "dashboard_04_priority_challenge",
        "dashboard_05_findings_report_generator",
    }
    missing = required - set(dashboards)
    if missing:
        raise ValueError(f"Report inputs are incomplete; missing dashboards: {sorted(missing)}")
    return dashboards


def _numeric_value(row: dict[str, Any]) -> float | None:
    """Return the first report-safe aggregate monetary value in a row."""
    for key in ("amount", "funding_amount", "above_baseline", "affected_dollars", "materiality"):
        value = row.get(key)
        if isinstance(value, (int, float)):
            return float(value)
    return None


def _label_value(row: dict[str, Any]) -> str:
    """Return a concise aggregate label for a report exhibit."""
    for key in (
        "fiscal_year",
        "program_request",
        "site",
        "organization",
        "acquisition_type",
        "title",
        "rule_id",
    ):
        value = row.get(key)
        if value not in (None, ""):
            return str(value)
    return "Aggregate"


def _finding_ids(payload: DashboardQuestionPayload) -> tuple[str, ...]:
    """Return bounded finding/rule identifiers represented by one aggregate payload."""
    return tuple(
        dict.fromkeys(
            str(identifier)
            for row in payload.data
            if (identifier := row.get("finding_id") or row.get("rule_id"))
        )
    )


def render_exhibit_png(payload: DashboardQuestionPayload, output: Path) -> Path:
    """Render a simple offline PNG from the same validated aggregate payload as the dashboard."""
    rows = [row for row in payload.data if _numeric_value(row) is not None][:8]
    width, height = 1200, max(360, 120 + 56 * max(len(rows), 1))
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((40, 24), payload.chart_title, fill="#142433")
    draw.text((40, 52), payload.question_text, fill="#425466")
    if not rows:
        draw.text(
            (40, 110), "No numeric aggregate values are available for this exhibit.", fill="#6a7785"
        )
    else:
        maximum = max(abs(_numeric_value(row) or 0.0) for row in rows) or 1.0
        for index, row in enumerate(rows):
            y = 100 + index * 56
            value = _numeric_value(row) or 0.0
            label = _label_value(row)[:42]
            draw.text((40, y + 8), label, fill="#142433")
            bar_start = 350
            bar_width = int(600 * abs(value) / maximum)
            color = "#8f3145" if value < 0 else "#1f5d78"
            draw.rectangle((bar_start, y, bar_start + max(bar_width, 2), y + 30), fill=color)
            draw.text((970, y + 8), format_dollars(value), fill="#142433")
    output.parent.mkdir(parents=True, exist_ok=True)
    temporary = output.with_suffix(".tmp.png")
    image.save(temporary, format="PNG")
    os.replace(temporary, output)
    return output


def _add_claim(
    document: DocumentType,
    citations: list[CitationRecord],
    *,
    paragraph_id: str,
    text: str,
    payload: DashboardQuestionPayload,
) -> None:
    """Add one deterministic claim and its payload citation."""
    document.add_paragraph(text)
    citations.append(
        CitationRecord(
            report_id=REPORT_ID,
            paragraph_id=paragraph_id,
            claim_text=text,
            citation_type="dashboard_payload",
            citation_id=payload.chart_id,
            source_file=payload.source_file,
            source_filter=payload.active_filter_state,
            chart_id=payload.chart_id,
            dashboard_id=payload.dashboard_id,
            metric_definition=payload.metric_definition,
            finding_ids=_finding_ids(payload),
            retrieval_chunk_id=f"rag:{payload.chart_id}",
            ontology_path=payload.ontology_references,
        )
    )


def _write_html_report(path: Path, sections: list[tuple[str, list[str]]]) -> None:
    """Write a Word-independent HTML companion using deterministic text only."""
    body = "\n".join(
        f"<section><h2>{html.escape(title)}</h2>"
        + "".join(f"<p>{html.escape(paragraph)}</p>" for paragraph in paragraphs)
        + "</section>"
        for title, paragraphs in sections
    )
    content = (
        "<!doctype html><html lang='en'><head><meta charset='utf-8'>"
        "<title>CEPE Pit Production FYNSP Program Review</title></head><body>"
        "<h1>CEPE Pit Production FYNSP Program Review</h1>"
        "<p>Deterministic draft; analyst review required.</p>"
        f"{body}</body></html>"
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(content, encoding="utf-8")
    os.replace(temporary, path)


def generate_management_report(project_root: Path | None = None) -> dict[str, Any]:
    """Generate DOCX, HTML, exhibits, report manifest, and citation manifest."""
    root = (project_root or Path(__file__).resolve().parents[3]).resolve()
    settings = load_settings(project_root=root)
    dashboards = _load_payloads(root)
    d1 = dashboards["dashboard_01_pit_production"]
    d2 = dashboards["dashboard_02_acquisition_schedule"]
    d3 = dashboards["dashboard_03_site_capacity"]
    d4 = dashboards["dashboard_04_priority_challenge"]
    d5 = dashboards["dashboard_05_findings_report_generator"]
    report_root = settings.resolve_path(settings.paths.reports_dir)
    docx_path = report_root / settings.report.docx_subdir / f"{REPORT_ID}.docx"
    html_path = report_root / settings.report.html_subdir / f"{REPORT_ID}.html"
    exhibit_root = report_root / settings.report.exhibit_subdir
    exhibits = {
        "funding": render_exhibit_png(d1["q1"], exhibit_root / "exhibit_01_funding.png"),
        "sites": render_exhibit_png(d3["q1"], exhibit_root / "exhibit_02_sites.png"),
        "requests": render_exhibit_png(d4["q1"], exhibit_root / "exhibit_03_requests.png"),
        "tier": render_exhibit_png(d4["q2"], exhibit_root / "exhibit_04_tier_flags.png"),
        "acquisition": render_exhibit_png(
            d2["q4"], exhibit_root / "exhibit_05_acquisition_completeness.png"
        ),
        "quality": render_exhibit_png(
            d1["q5"], exhibit_root / "exhibit_06_quality_traceability.png"
        ),
        "reconciliation": render_exhibit_png(
            d1["q6"], exhibit_root / "exhibit_07_reconciliation.png"
        ),
    }
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    title = document.add_heading(settings.report.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        "Deterministic management-review draft generated from validated aggregate dashboard evidence. "
        "Analyst review is required before release."
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    citations: list[CitationRecord] = []
    section_text: list[tuple[str, list[str]]] = []

    section_payloads = [
        ("Executive Summary", [d1["q1"], d1["q6"], d5["q1"]]),
        ("Review Scope and Data Sources", [d1["q1"]]),
        (
            "Dashboard Question Answers",
            [payload for dashboard in dashboards.values() for payload in dashboard.values()],
        ),
        ("Accuracy Findings", [d1["q5"], d1["q6"], d5["q1"]]),
        ("Thoroughness Findings", [d5["q2"], d4["q5"], d4["q6"]]),
        ("Uncertainty, Risk, and Opportunity Findings", [d5["q3"], d3["q5"], d2["q4"]]),
        ("Data Limitations and Recommended Follow-Up", [d1["q1"], d5["q1"], d5["q5"]]),
        ("Appendix: Source Lineage and Guidance Citations", [d5["q5"], d5["q6"]]),
    ]
    paragraph_index = 0
    for section_index, (section, payloads) in enumerate(section_payloads):
        document.add_heading(section, level=1)
        paragraphs: list[str] = []
        if section == "Review Scope and Data Sources":
            document.add_heading("Scope and Methodology", level=2)
            paragraphs.append(
                "The review covers FY2028-FY2032 FORMEX Pit Production programming. Federal Crosscuts "
                "support portfolio totals and Federal Site Splits support site analysis. PLANEX and "
                "COSTEX remain execution context and are not directly reconciled without a crosswalk."
            )
            document.add_paragraph(paragraphs[-1])
        if section == "Dashboard Question Answers":
            document.add_heading("Programmed Funding Overview", level=2)
            document.add_heading("Acquisition and Schedule Observations", level=2)
            table = document.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for cell, heading in zip(
                table.rows[0].cells,
                ("Dashboard", "Question", "Calculated answer", "Evidence ID"),
                strict=True,
            ):
                cell.text = heading
            for payload in payloads:
                claim = payload.plain_language_summary
                cells = table.add_row().cells
                cells[0].text = payload.dashboard_title
                cells[1].text = payload.question_text
                cells[2].text = claim
                finding_ids = _finding_ids(payload)
                cells[3].text = ", ".join((payload.chart_id, *finding_ids))
                paragraph_index += 1
                citations.append(
                    CitationRecord(
                        report_id=REPORT_ID,
                        paragraph_id=f"p{paragraph_index:03d}",
                        claim_text=claim,
                        citation_type="dashboard_payload",
                        citation_id=payload.chart_id,
                        source_file=payload.source_file,
                        source_filter=payload.active_filter_state,
                        chart_id=payload.chart_id,
                        dashboard_id=payload.dashboard_id,
                        metric_definition=payload.metric_definition,
                        finding_ids=finding_ids,
                        retrieval_chunk_id=f"rag:{payload.chart_id}",
                        ontology_path=payload.ontology_references,
                    )
                )
                paragraphs.append(f"{payload.question_text} — {claim}")
            document.add_heading("Recommended exhibits", level=2)
            exhibit_table = document.add_table(rows=0, cols=2)
            for index, (exhibit_id, image) in enumerate(exhibits.items()):
                if index % 2 == 0:
                    exhibit_cells = exhibit_table.add_row().cells
                cell = exhibit_cells[index % 2]
                cell.paragraphs[0].add_run().add_picture(str(image), width=Inches(2.85))
                cell.add_paragraph(exhibit_id.replace("_", " ").title())
            section_text.append((section, paragraphs))
            document.add_page_break()
            continue
        if section == "Accuracy Findings":
            document.add_heading("Data-Quality and Reconciliation Findings", level=2)
        elif section == "Uncertainty, Risk, and Opportunity Findings":
            document.add_heading("Site-Capacity Observations", level=2)
            document.add_heading("Priority and Above-Baseline Observations", level=2)
        elif section == "Data Limitations and Recommended Follow-Up":
            document.add_heading("Management Actions and Recommendations", level=2)
            document.add_heading("Limitations", level=2)
        elif section == "Appendix: Source Lineage and Guidance Citations":
            document.add_heading("Citation and Evidence Manifest", level=2)
        for payload in payloads:
            paragraph_index += 1
            claim = payload.plain_language_summary
            if section == "Data Limitations and Recommended Follow-Up":
                claim = (
                    "Analyst action: review the evidence and limitations associated with "
                    f"{payload.chart_id}; the dashboard does not authorize write-back or operational action."
                )
            _add_claim(
                document,
                citations,
                paragraph_id=f"p{paragraph_index:03d}",
                text=claim,
                payload=payload,
            )
            paragraphs.append(claim)
        section_text.append((section, paragraphs))
        if section_index in {0, 5}:
            document.add_page_break()

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    temporary_docx = docx_path.with_suffix(".tmp.docx")
    document.save(str(temporary_docx))
    os.replace(temporary_docx, docx_path)
    _write_html_report(html_path, section_text)
    citation_path = report_root / "citation_manifest.json"
    citation_payload = {
        "schema_version": "2.0",
        "report_id": REPORT_ID,
        "citations": [record.model_dump(mode="json") for record in citations],
    }
    write_json(citation_path, citation_payload)
    manifest = {
        "schema_version": "2.0",
        "report_id": REPORT_ID,
        "status": "deterministic_draft_generated",
        "target_length": f"{settings.report.target_pages} pages under normal Word rendering",
        "sections": list(REPORT_SECTIONS),
        "docx_file": docx_path.relative_to(root).as_posix(),
        "html_file": html_path.relative_to(root).as_posix(),
        "citation_manifest_file": citation_path.relative_to(root).as_posix(),
        "exhibits": [path.relative_to(root).as_posix() for path in exhibits.values()],
        "exhibit_sources": {
            "exhibit_01_funding": d1["q1"].chart_id,
            "exhibit_02_sites": d3["q1"].chart_id,
            "exhibit_03_requests": d4["q1"].chart_id,
            "exhibit_04_tier_flags": d4["q2"].chart_id,
            "exhibit_05_acquisition_completeness": d2["q4"].chart_id,
            "exhibit_06_quality_traceability": d1["q5"].chart_id,
            "exhibit_07_reconciliation": d1["q6"].chart_id,
        },
        "finding_ids": list(_finding_ids(d5["q1"])),
        "citation_count": len(citations),
        "narrative_origin": "deterministic_analytical_conclusion",
        "ai_assisted_text": False,
        "limitations": [
            "No guidance documents were available for guidance-passage citations.",
            "The report is a deterministic draft and requires analyst review.",
            "PLANEX/COSTEX are not directly reconciled to FY2028-FY2032 FORMEX without a crosswalk.",
        ],
    }
    write_json(
        report_root / settings.report.html_subdir / "dashboard_05_report_manifest.json", manifest
    )
    return manifest
